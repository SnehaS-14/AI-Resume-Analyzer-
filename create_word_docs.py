from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_style(table):
    """Add styling to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

def create_api_documentation():
    """Create API Documentation Word file"""
    doc = Document()

    # Title
    title = doc.add_heading('API DOCUMENTATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(16, 185, 129)

    # Subtitle
    subtitle = doc.add_paragraph('AI Resume Analyzer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True

    doc.add_paragraph('Base URL: https://ai-resume-analyzer-svry.onrender.com')
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        'Authentication Endpoints',
        'Resume Analysis Endpoints',
        'History Endpoints',
        'Rewrite Endpoints',
        'Error Handling',
        'Data Models'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # Overview
    doc.add_heading('Overview', 1)
    doc.add_paragraph('The AI Resume Analyzer API is a FastAPI-based REST API that provides endpoints for:')
    doc.add_paragraph('User authentication (signup & login)', style='List Bullet')
    doc.add_paragraph('Resume analysis using AI', style='List Bullet')
    doc.add_paragraph('Resume history management', style='List Bullet')
    doc.add_paragraph('Resume rewriting capabilities', style='List Bullet')

    doc.add_page_break()

    # Authentication Endpoints
    doc.add_heading('Authentication Endpoints', 1)

    # Signup
    doc.add_heading('1. User Signup', 2)
    doc.add_paragraph('POST /auth/signup')
    doc.add_paragraph('Register a new user account')

    doc.add_heading('Request Body:', 3)
    doc.add_paragraph('{')
    doc.add_paragraph('"name": "John Doe",', style='List Bullet')
    doc.add_paragraph('"email": "john@example.com",', style='List Bullet')
    doc.add_paragraph('"password": "securePassword123"', style='List Bullet')
    doc.add_paragraph('}')

    # Parameters table
    doc.add_heading('Request Parameters:', 3)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'

    data = [
        ('name', 'string', 'Full name of the user'),
        ('email', 'string', 'Email address (must be valid)'),
        ('password', 'string', 'Password (min 6 characters)')
    ]

    for i, (param, ptype, desc) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = param
        row_cells[1].text = ptype
        row_cells[2].text = desc

    doc.add_heading('Response (200 OK):', 3)
    doc.add_paragraph('{')
    doc.add_paragraph('"user_id": "507f1f77bcf86cd799439011",', style='List Bullet')
    doc.add_paragraph('"email": "john@example.com",', style='List Bullet')
    doc.add_paragraph('"name": "John Doe",', style='List Bullet')
    doc.add_paragraph('"message": "Signup successful"', style='List Bullet')
    doc.add_paragraph('}')

    doc.add_heading('Errors:', 3)
    err_table = doc.add_table(rows=3, cols=2)
    err_table.style = 'Light Grid Accent 1'
    err_hdr = err_table.rows[0].cells
    err_hdr[0].text = 'Status'
    err_hdr[1].text = 'Message'

    errors = [
        ('400', 'Email already registered'),
        ('422', 'Invalid email format or missing fields')
    ]

    for i, (status, msg) in enumerate(errors, 1):
        row = err_table.rows[i].cells
        row[0].text = status
        row[1].text = msg

    # Login
    doc.add_page_break()
    doc.add_heading('2. User Login', 2)
    doc.add_paragraph('POST /auth/login')
    doc.add_paragraph('Authenticate user and retrieve credentials')

    doc.add_heading('Request Body:', 3)
    doc.add_paragraph('{')
    doc.add_paragraph('"email": "john@example.com",', style='List Bullet')
    doc.add_paragraph('"password": "securePassword123"', style='List Bullet')
    doc.add_paragraph('}')

    doc.add_heading('Response (200 OK):', 3)
    doc.add_paragraph('{')
    doc.add_paragraph('"user_id": "507f1f77bcf86cd799439011",', style='List Bullet')
    doc.add_paragraph('"email": "john@example.com",', style='List Bullet')
    doc.add_paragraph('"name": "John Doe",', style='List Bullet')
    doc.add_paragraph('"message": "Login successful"', style='List Bullet')
    doc.add_paragraph('}')

    # Resume Analysis
    doc.add_page_break()
    doc.add_heading('Resume Analysis Endpoints', 1)

    doc.add_heading('3. Analyze Resume', 2)
    doc.add_paragraph('POST /analyze')
    doc.add_paragraph('Upload and analyze a resume using AI')

    doc.add_heading('Query Parameters:', 3)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'
    row = table.rows[1].cells
    row[0].text = 'user_id'
    row[1].text = 'string'
    row[2].text = 'MongoDB ObjectId of the user'

    doc.add_heading('Response Fields:', 3)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    fields = [
        ('overall_score', 'integer', 'Resume score 0-100'),
        ('ats_score', 'integer', 'ATS compatibility 0-100'),
        ('strengths', 'array', 'Resume strengths'),
        ('weaknesses', 'array', 'Areas for improvement'),
        ('action_items', 'array', 'Specific recommendations'),
        ('summary', 'string', 'AI-generated summary')
    ]

    for i, (field, ftype, desc) in enumerate(fields, 1):
        row = table.rows[i].cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    # History
    doc.add_page_break()
    doc.add_heading('History Endpoints', 1)

    doc.add_heading('4. Get User Analysis History', 2)
    doc.add_paragraph('GET /history')
    doc.add_paragraph('Retrieve list of past resume analyses for a user')

    doc.add_heading('Query Parameters:', 3)
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Type'
    hdr[2].text = 'Default'
    hdr[3].text = 'Description'

    params = [
        ('user_id', 'string', '-', 'User MongoDB ID'),
        ('limit', 'integer', '20', 'Max records to return')
    ]

    for i, (param, ptype, default, desc) in enumerate(params, 1):
        row = table.rows[i].cells
        row[0].text = param
        row[1].text = ptype
        row[2].text = default
        row[3].text = desc

    doc.add_heading('5. Get Analysis Details', 2)
    doc.add_paragraph('GET /history/{id}')
    doc.add_paragraph('Get full details of a specific resume analysis')

    # Rewrite
    doc.add_page_break()
    doc.add_heading('Rewrite Endpoints', 1)

    doc.add_heading('6. Rewrite Resume', 2)
    doc.add_paragraph('POST /rewrite/{id}')
    doc.add_paragraph('Generate an improved version of a resume')

    doc.add_heading('Path Parameters:', 3)
    doc.add_paragraph('id - MongoDB ObjectId of the analysis', style='List Bullet')

    doc.add_heading('Query Parameters:', 3)
    doc.add_paragraph('user_id - User MongoDB ID (for authorization)', style='List Bullet')

    doc.add_heading('Response (200 OK):', 3)
    doc.add_paragraph('{')
    doc.add_paragraph('"rewritten_text": "Improved resume text..."', style='List Bullet')
    doc.add_paragraph('}')

    doc.add_heading('7. Download Rewritten Resume', 2)
    doc.add_paragraph('GET /rewrite/{id}/download')
    doc.add_paragraph('Download the rewritten resume as a .txt file')

    # Error Handling
    doc.add_page_break()
    doc.add_heading('Error Handling', 1)

    doc.add_heading('Error Response Format:', 2)
    doc.add_paragraph('{')
    doc.add_paragraph('"detail": "Description of the error"', style='List Bullet')
    doc.add_paragraph('}')

    doc.add_heading('HTTP Status Codes:', 2)
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Status Code'
    hdr[1].text = 'Meaning'

    codes = [
        ('200', 'Success'),
        ('400', 'Bad Request'),
        ('401', 'Unauthorized'),
        ('403', 'Forbidden'),
        ('404', 'Not Found'),
        ('413', 'Payload Too Large'),
        ('422', 'Unprocessable Entity'),
        ('500', 'Internal Server Error')
    ]

    for i, (code, meaning) in enumerate(codes, 1):
        row = table.rows[i].cells
        row[0].text = code
        row[1].text = meaning

    # Data Models
    doc.add_page_break()
    doc.add_heading('Data Models', 1)

    doc.add_heading('User Model', 2)
    doc.add_paragraph('_id: ObjectId - MongoDB ID', style='List Bullet')
    doc.add_paragraph('email: string - User email (unique, lowercase)', style='List Bullet')
    doc.add_paragraph('name: string - User full name', style='List Bullet')
    doc.add_paragraph('password_hash: string - PBKDF2 hashed password', style='List Bullet')
    doc.add_paragraph('created_at: datetime - Account creation timestamp', style='List Bullet')

    doc.add_heading('Resume Analysis Model', 2)
    doc.add_paragraph('_id: ObjectId - MongoDB ID', style='List Bullet')
    doc.add_paragraph('user_id: string - User MongoDB ID', style='List Bullet')
    doc.add_paragraph('filename: string - Original file name', style='List Bullet')
    doc.add_paragraph('uploaded_at: datetime - Upload timestamp', style='List Bullet')
    doc.add_paragraph('overall_score: integer - Overall score (0-100)', style='List Bullet')
    doc.add_paragraph('ats_score: integer - ATS score (0-100)', style='List Bullet')
    doc.add_paragraph('strengths: array - List of strengths', style='List Bullet')
    doc.add_paragraph('weaknesses: array - List of weaknesses', style='List Bullet')
    doc.add_paragraph('action_items: array - Improvement recommendations', style='List Bullet')
    doc.add_paragraph('summary: string - AI-generated summary', style='List Bullet')
    doc.add_paragraph('raw_text: string - Extracted resume text', style='List Bullet')
    doc.add_paragraph('rewritten_text: string - AI-rewritten resume (nullable)', style='List Bullet')

    # Save
    doc.save('API_DOCUMENTATION.docx')
    print("[OK] API_DOCUMENTATION.docx created successfully!")

def create_project_flow_documentation():
    """Create Project Flow Documentation Word file"""
    doc = Document()

    # Title
    title = doc.add_heading('PROJECT FLOW & ARCHITECTURE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(16, 185, 129)

    # Subtitle
    subtitle = doc.add_paragraph('AI Resume Analyzer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True

    doc.add_paragraph()

    # Overview
    doc.add_heading('Project Overview', 1)
    doc.add_paragraph('AI Resume Analyzer is a web application that helps job seekers improve their resumes using AI-powered analysis.')

    doc.add_heading('Key Features:', 2)
    doc.add_paragraph('Resume upload and analysis', style='List Bullet')
    doc.add_paragraph('AI-driven feedback with scores', style='List Bullet')
    doc.add_paragraph('Action items for improvement', style='List Bullet')
    doc.add_paragraph('Resume rewriting with AI', style='List Bullet')
    doc.add_paragraph('Analysis history tracking', style='List Bullet')
    doc.add_paragraph('Download rewritten resumes', style='List Bullet')

    doc.add_page_break()

    # Technology Stack
    doc.add_heading('Technology Stack', 1)

    doc.add_heading('Frontend', 2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Technology'
    hdr[2].text = 'Purpose'

    frontend_tech = [
        ('Framework', 'React 19.2', 'UI rendering'),
        ('Language', 'TypeScript 6.0', 'Type safety'),
        ('Build Tool', 'Vite 8.0', 'Fast bundling'),
        ('Styling', 'Inline CSS + Tailwind', 'Responsive design'),
        ('HTTP Client', 'Axios 1.16', 'API requests'),
        ('Deployment', 'Render/Vercel', 'Static hosting')
    ]

    for i, (comp, tech, purpose) in enumerate(frontend_tech, 1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = purpose

    doc.add_heading('Backend', 2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Technology'
    hdr[2].text = 'Purpose'

    backend_tech = [
        ('Framework', 'FastAPI', 'REST API server'),
        ('Language', 'Python 3.11', 'Server-side logic'),
        ('Server', 'Uvicorn', 'ASGI server'),
        ('Database', 'MongoDB', 'NoSQL document store'),
        ('Driver', 'Motor', 'Async MongoDB driver'),
        ('AI API', 'Groq', 'LLM-powered analysis')
    ]

    for i, (comp, tech, purpose) in enumerate(backend_tech, 1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = tech
        row[2].text = purpose

    doc.add_page_break()

    # Data Flow
    doc.add_heading('User Registration Flow', 1)
    doc.add_paragraph('User Input (Email, Password, Name)', style='List Number')
    doc.add_paragraph('Frontend validation', style='List Number')
    doc.add_paragraph('POST /auth/signup', style='List Number')
    doc.add_paragraph('Backend: Hash password (PBKDF2)', style='List Number')
    doc.add_paragraph('Check if email exists', style='List Number')
    doc.add_paragraph('Insert user to MongoDB', style='List Number')
    doc.add_paragraph('Response: user_id, email, name', style='List Number')
    doc.add_paragraph('Frontend: Save to localStorage', style='List Number')
    doc.add_paragraph('Redirect to main app', style='List Number')

    doc.add_page_break()

    doc.add_heading('Resume Analysis Flow', 1)
    doc.add_paragraph('User uploads resume file', style='List Number')
    doc.add_paragraph('Frontend sends file + user_id', style='List Number')
    doc.add_paragraph('POST /analyze?user_id=XXX', style='List Number')
    doc.add_paragraph('Backend receives file', style='List Number')
    doc.add_paragraph('Extract text using:', style='List Number')
    doc.add_paragraph('pdfplumber (for PDF)', style='List Bullet 2')
    doc.add_paragraph('python-docx (for DOCX)', style='List Bullet 2')
    doc.add_paragraph('Direct read (for TXT)', style='List Bullet 2')
    doc.add_paragraph('Send to Groq AI API for analysis', style='List Number')
    doc.add_paragraph('Groq returns:', style='List Number')
    doc.add_paragraph('overall_score, ats_score', style='List Bullet 2')
    doc.add_paragraph('strengths, weaknesses', style='List Bullet 2')
    doc.add_paragraph('action_items, summary', style='List Bullet 2')
    doc.add_paragraph('Save analysis to MongoDB', style='List Number')
    doc.add_paragraph('Return JSON response with document ID', style='List Number')
    doc.add_paragraph('Frontend displays results', style='List Number')

    doc.add_page_break()

    doc.add_heading('Resume Rewrite Flow', 1)
    doc.add_paragraph('User clicks "Rewrite Resume"', style='List Number')
    doc.add_paragraph('Frontend: POST /rewrite/{id}?user_id=XXX', style='List Number')
    doc.add_paragraph('Backend fetches analysis from DB', style='List Number')
    doc.add_paragraph('Extract raw_text and action_items', style='List Number')
    doc.add_paragraph('Send to Groq API with improvement instructions', style='List Number')
    doc.add_paragraph('Groq returns improved resume text', style='List Number')
    doc.add_paragraph('Update MongoDB with rewritten_text', style='List Number')
    doc.add_paragraph('Response: rewritten resume text', style='List Number')
    doc.add_paragraph('Frontend displays improved version', style='List Number')
    doc.add_paragraph('User can download as .txt file', style='List Number')

    doc.add_page_break()

    # Deployment
    doc.add_heading('Deployment Architecture', 1)

    doc.add_heading('Production Environment (Render)', 2)
    doc.add_paragraph('Frontend Service (Static Site)', style='List Number')
    doc.add_paragraph('Name: resume-ai-frontend', style='List Bullet 2')
    doc.add_paragraph('Framework: Vite (Node)', style='List Bullet 2')
    doc.add_paragraph('Build: npm install && npm run build', style='List Bullet 2')
    doc.add_paragraph('URL: https://ai-resume-analyzer-1-5veb.onrender.com', style='List Bullet 2')

    doc.add_paragraph('Backend Service (Web Service)', style='List Number')
    doc.add_paragraph('Name: resume-ai-backend', style='List Bullet 2')
    doc.add_paragraph('Language: Python 3.11', style='List Bullet 2')
    doc.add_paragraph('Build: pip install -r requirements.txt', style='List Bullet 2')
    doc.add_paragraph('Start: uvicorn main:app --host 0.0.0.0 --port $PORT', style='List Bullet 2')
    doc.add_paragraph('URL: https://ai-resume-analyzer-svry.onrender.com', style='List Bullet 2')

    doc.add_paragraph('External Services:', style='List Number')
    doc.add_paragraph('MongoDB Atlas (Cloud Database)', style='List Bullet 2')
    doc.add_paragraph('Groq API (LLM Service)', style='List Bullet 2')

    doc.add_page_break()

    # Database Schema
    doc.add_heading('Database Schema', 1)

    doc.add_heading('Users Collection', 2)
    doc.add_paragraph('_id: ObjectId', style='List Bullet')
    doc.add_paragraph('email: string (unique, lowercase)', style='List Bullet')
    doc.add_paragraph('name: string', style='List Bullet')
    doc.add_paragraph('password_hash: string', style='List Bullet')
    doc.add_paragraph('created_at: Date', style='List Bullet')

    doc.add_paragraph('Indexes: email (unique), created_at', style='List Bullet')

    doc.add_heading('Analyses Collection', 2)
    doc.add_paragraph('_id: ObjectId', style='List Bullet')
    doc.add_paragraph('user_id: string (ObjectId)', style='List Bullet')
    doc.add_paragraph('filename: string', style='List Bullet')
    doc.add_paragraph('uploaded_at: Date', style='List Bullet')
    doc.add_paragraph('overall_score: number (0-100)', style='List Bullet')
    doc.add_paragraph('ats_score: number (0-100)', style='List Bullet')
    doc.add_paragraph('strengths: [string]', style='List Bullet')
    doc.add_paragraph('weaknesses: [string]', style='List Bullet')
    doc.add_paragraph('action_items: [string]', style='List Bullet')
    doc.add_paragraph('summary: string', style='List Bullet')
    doc.add_paragraph('raw_text: string', style='List Bullet')
    doc.add_paragraph('rewritten_text: string (optional)', style='List Bullet')

    doc.add_paragraph('Indexes: user_id, uploaded_at, (user_id, uploaded_at)', style='List Bullet')

    doc.add_page_break()

    # Security
    doc.add_heading('Security Measures', 1)
    doc.add_paragraph('Password Security', style='List Number')
    doc.add_paragraph('PBKDF2-HMAC-SHA256 hashing', style='List Bullet 2')
    doc.add_paragraph('100,000 iterations', style='List Bullet 2')
    doc.add_paragraph('Random salt (32 bytes)', style='List Bullet 2')

    doc.add_paragraph('CORS Protection', style='List Number')
    doc.add_paragraph('Whitelist specific origins', style='List Bullet 2')
    doc.add_paragraph('Allow only necessary methods', style='List Bullet 2')
    doc.add_paragraph('Validate headers', style='List Bullet 2')

    doc.add_paragraph('Input Validation', style='List Number')
    doc.add_paragraph('Email validation (EmailStr)', style='List Bullet 2')
    doc.add_paragraph('File type checking', style='List Bullet 2')
    doc.add_paragraph('File size limits (10MB)', style='List Bullet 2')

    doc.add_paragraph('Authorization', style='List Number')
    doc.add_paragraph('user_id parameter validation', style='List Bullet 2')
    doc.add_paragraph('Users can only access their own data', style='List Bullet 2')
    doc.add_paragraph('MongoDB ObjectId validation', style='List Bullet 2')

    doc.add_paragraph('Database Security', style='List Number')
    doc.add_paragraph('MongoDB Atlas IP whitelist', style='List Bullet 2')
    doc.add_paragraph('Connection string in environment variables', style='List Bullet 2')
    doc.add_paragraph('No sensitive data in logs', style='List Bullet 2')

    # Save
    doc.save('PROJECT_FLOW.docx')
    print("✓ PROJECT_FLOW.docx created successfully!")

# Create both documents
if __name__ == '__main__':
    try:
        create_api_documentation()
        create_project_flow_documentation()
        print("\n✅ Word documents created successfully!")
        print("Files created:")
        print("  - API_DOCUMENTATION.docx")
        print("  - PROJECT_FLOW.docx")
    except Exception as e:
        print(f"Error: {e}")
